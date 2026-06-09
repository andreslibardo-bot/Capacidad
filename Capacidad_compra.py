import pandas as pd
import numpy as np
import re
import os
from fractions import Fraction
from datetime import datetime
import warnings
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

warnings.simplefilter(action="ignore", category=DeprecationWarning)
warnings.filterwarnings("ignore", category=FutureWarning)
pd.set_option("display.max_rows", None)
pd.set_option("display.max_columns", None)
pd.set_option('display.float_format', '{:,.2f}'.format)

warnings.simplefilter(action="ignore", category=DeprecationWarning)
warnings.filterwarnings("ignore", category=FutureWarning)
pd.set_option("display.max_rows", None)
pd.set_option("display.max_columns", None)


############################################  Funciones Corrida de capacidad cilindros y tanques y redes¶ ##########################################

#################################################  funcion General #################################################################################


#####################   Funcion para calcular la capacidad de compra de cilindros y tanques ########################################################

def calcular_CCit(data_cil, data_tanques_T2, data_tanques_T1):
    # Normalizar nombres de columnas (equivalente a make.names en R)
    data_cil.columns = [re.sub(r'\W+', '_', col) for col in data_cil.columns]
    
    # Convertir tipos
    data_cil['ID_EMPRESA'] = data_cil['ID_EMPRESA'].astype(int)
    data_cil['EMPRESA'] = data_cil['EMPRESA'].astype(str)
    data_cil['ID_MARCA'] = data_cil['ID_MARCA'].astype(str)
    data_cil['MARCA'] = data_cil['MARCA'].astype(str)
    data_cil['ORIGEN'] = data_cil['ORIGEN'].astype(str)
    data_cil['CAPACIDAD'] = data_cil['CAPACIDAD'].str.upper()
    data_cil['CONTEO'] = data_cil['CONTEO'].astype(int)
    
    # Extraer unidades y masa
    data_cil['UNIDADES'] = data_cil['CAPACIDAD'].str.extract(r'(?<=\s)(.*)')[0]
    data_cil['MASA'] = data_cil['CAPACIDAD'].str.extract(r'^(\d+)')[0].astype(int)
    
    # Calcular capacidad en kg
    data_cil['Capacidad_Cil_kg'] = np.where(  data_cil['UNIDADES'] == 'LIBRAS',   data_cil['MASA'] * data_cil['CONTEO'] * 0.454 * 6,  data_cil['MASA'] * data_cil['CONTEO'] * 6 )
    
    data_cil = data_cil.sort_values(['ID_EMPRESA', 'MARCA'])
    
    # Agrupar por empresa
    data_Cilindros_CapCil = ( data_cil.groupby(['ID_EMPRESA', 'EMPRESA'], as_index=False)
        .agg(
            N_de_cilindros=('CONTEO', 'sum'),
            Cap_cil=('Capacidad_Cil_kg', 'sum')
        )
        .sort_values('ID_EMPRESA')
    )
    
    # --- Procesamiento de tanques T2 ---
    data_tanques_T2.columns = [re.sub(r'\W+', '_', col) for col in data_tanques_T2.columns]
    data_tanques_T2.rename(columns={'IDENTIFICADOR_EMPRESA':'ID_EMPRESA', 'ARE_ESP_NOMBRE':'EMPRESA'}, inplace=True)
    
    data_tanques_T2['ID_EMPRESA'] = data_tanques_T2['ID_EMPRESA'].astype(int)
    data_tanques_T2['CAR_T281_CAPACI'] = data_tanques_T2['CAR_T281_CAPACI'].astype(int)
    data_tanques_T2['CapTE_T2'] = data_tanques_T2['CAR_T281_CAPACI'] * 2.1 * 6
    
    data_tanques_T2_grup = (
        data_tanques_T2.groupby(['ID_EMPRESA', 'EMPRESA'], as_index=False)
        .agg(CapTE_T2=('CapTE_T2', 'sum'))
    )
    
    # --- Procesamiento de tanques T1 ---
    data_tanques_T1.columns = [re.sub(r'\W+', '_', col) for col in data_tanques_T1.columns]
    data_tanques_T1.rename(columns={'IDENTIFICADOR_EMPRESA':'ID_EMPRESA', 'ARE_ESP_NOMBRE':'EMPRESA'}, inplace=True)
    
    data_tanques_T1['ID_EMPRESA'] = data_tanques_T1['ID_EMPRESA'].astype(int)
    data_tanques_T1['CAR_T281_CAPACI'] = data_tanques_T1['CAR_T281_CAPACI'].astype(int)
    data_tanques_T1['CapTE_T1'] = data_tanques_T1['CAR_T281_CAPACI'] * 2.1 * 6
    
    data_tanques_T1_grup = (
        data_tanques_T1.groupby(['ID_EMPRESA', 'EMPRESA'], as_index=False)
        .agg(CapTE_T1=('CapTE_T1', 'sum'))
    )
    
    # --- Consolidado tanques ---
    data_Tanques_CapTE = pd.merge(data_tanques_T2_grup, data_tanques_T1_grup, on='ID_EMPRESA', how='outer', suffixes=('.x', '.y'))
    data_Tanques_CapTE['EMPRESA'] = data_Tanques_CapTE['EMPRESA.x'].combine_first(data_Tanques_CapTE['EMPRESA.y'])
    data_Tanques_CapTE['CapTE'] = np.where(
        (data_Tanques_CapTE['CapTE_T2'].isna()) | (data_Tanques_CapTE['CapTE_T2']==0),
        data_Tanques_CapTE['CapTE_T1'],
        data_Tanques_CapTE['CapTE_T2']
    )
    data_Tanques_CapTE = data_Tanques_CapTE[['ID_EMPRESA', 'EMPRESA', 'CapTE']]
    
    # --- Unión final ---
    merged_data = pd.merge(data_Tanques_CapTE, data_Cilindros_CapCil, on='ID_EMPRESA', how='outer', suffixes=('.x', '.y'))
    merged_data['EMPRESA'] = merged_data['EMPRESA.x'].combine_first(merged_data['EMPRESA.y'])
    merged_data['CapTE'] = merged_data['CapTE'].fillna(0)
    merged_data['Cap_cil'] = merged_data['Cap_cil'].fillna(0)
    merged_data['CCit'] = (0.85 * merged_data['CapTE'] + merged_data['Cap_cil']) * 0.345
    
    merged_data = merged_data[['ID_EMPRESA', 'EMPRESA', 'CapTE', 'Cap_cil', 'CCit']]
    
    return {
        'data_Cil': data_cil,
        'data_Cilindros_CapCil': data_Cilindros_CapCil,
        'data_tanques_T2_grup': data_tanques_T2_grup,
        'data_tanques_T2': data_tanques_T2,
        'data_tanques_T1_grup': data_tanques_T1_grup,
        'data_tanques_T1': data_tanques_T1,
        'data_Tanques_CapTE': data_Tanques_CapTE,
        'resultado': merged_data
    }


#######################################################################################  funcion para procesar tanques #################

def procesar_tanques(data_tanques_T2, data_tanques_T1,  debug: bool = False) -> pd.DataFrame:
     
    # --- 1. Leer y procesar T2 ---
    data_tanques_T2 = data_tanques_T2.rename(columns={'IDENTIFICADOR_EMPRESA':'ID_EMPRESA','ARE_ESP_NOMBRE':'EMPRESA'})
    data_tanques_T2['ID_EMPRESA'] = data_tanques_T2['ID_EMPRESA'].astype(int)
    data_tanques_T2['Cap_gal_T2'] = data_tanques_T2['CAR_T281_CAPACI'].astype(int)
    data_tanques_T2['CapTE_T2'] = data_tanques_T2['CAR_T281_CAPACI'] * 2.1 * 6

    data_tanques_T2_clean = (  data_tanques_T2.groupby(['ID_EMPRESA','EMPRESA'], as_index=False)
        .agg(
            CapTE_T2=('CapTE_T2','sum'),
            n_T2=('CapTE_T2','size'),
            Cap_gal_T2=('Cap_gal_T2','sum')
        )
    )

    # --- 2. Leer y procesar T1 ---
    data_tanques_T1 = data_tanques_T1.rename(columns={'IDENTIFICADOR_EMPRESA':'ID_EMPRESA','ARE_ESP_NOMBRE':'EMPRESA'})
    data_tanques_T1['ID_EMPRESA'] = data_tanques_T1['ID_EMPRESA'].astype(int)
    data_tanques_T1['Cap_gal_T1'] = data_tanques_T1['CAR_T281_CAPACI'].astype(int)
    data_tanques_T1['CapTE_T1'] = data_tanques_T1['CAR_T281_CAPACI'] * 2.1 * 6

    data_tanques_T1_clean = (  data_tanques_T1.groupby(['ID_EMPRESA','EMPRESA'], as_index=False)
        .agg(
            CapTE_T1=('CapTE_T1','sum'),
            n_T1=('CapTE_T1','size'),
            Cap_gal_T1=('Cap_gal_T1','sum')
        )
    )

    # --- 3. Consolidar T2 y T1 ---
    data_Tanques_CapTE = pd.merge(  data_tanques_T2_clean,  data_tanques_T1_clean,  on='ID_EMPRESA', how='outer', suffixes=('.T2','.T1')  )

    data_Tanques_CapTE['Empresa'] = data_Tanques_CapTE['EMPRESA.T2'].combine_first(data_Tanques_CapTE['EMPRESA.T1'])
    data_Tanques_CapTE['CapTEi_t_kg'] = data_Tanques_CapTE['CapTE_T2'].combine_first(data_Tanques_CapTE['CapTE_T1'])
    data_Tanques_CapTE['Número_de_Tanques'] = data_Tanques_CapTE['n_T2'].combine_first(data_Tanques_CapTE['n_T1'])
    data_Tanques_CapTE['Capacidad_gal'] = data_Tanques_CapTE['Cap_gal_T2'].combine_first(data_Tanques_CapTE['Cap_gal_T1'])

    # Redondear columnas numéricas
    cols_numeric = ['CapTEi_t_kg','Número_de_Tanques','Capacidad_gal']
    data_Tanques_CapTE[cols_numeric] = data_Tanques_CapTE[cols_numeric].round(0)

    # Selección de columnas finales
    data_Tanques_CapTE_final = data_Tanques_CapTE[['ID_EMPRESA','Capacidad_gal','Número_de_Tanques','CapTEi_t_kg']]
    data_Tanques_CapTE_final = data_Tanques_CapTE_final.rename(columns={'ID_EMPRESA':'Código SUI'})
    data_Tanques_CapTE_final = data_Tanques_CapTE_final.sort_values('Código SUI').reset_index(drop=True)

    if debug:
        return {
            'T2_clean': data_tanques_T2_clean,
            'T1_clean': data_tanques_T1_clean,
            'Consolidado': data_Tanques_CapTE,
            'Final': data_Tanques_CapTE_final
        }

    return data_Tanques_CapTE_final


############################################################################   funcion para procesar cilindros ########################


def procesar_cilindros(Cilindros, debug: bool = False) -> pd.DataFrame:
    # --- 2. Separar número y unidad ---
    Cilindros[['Masa', 'Medida']] = Cilindros['CAPACIDAD'].str.split(" ", expand=True)
    Cilindros['Masa'] = pd.to_numeric(Cilindros['Masa'], errors='coerce')
    Cilindros['Medida'] = Cilindros['Medida'].str.upper()
    Cilindros['Medida'] = Cilindros['Medida'].replace({'LIBRAS':'LB', 'LB':'LB'})

    # --- 3. Cilindros en LB ---
    cilindros_lb = (  Cilindros[Cilindros["Medida"] == "LB"].groupby(["ID_EMPRESA", "EMPRESA", "Masa"], as_index=False)["CONTEO"].sum() )
    cilindros_lb = cilindros_lb.pivot(index=["ID_EMPRESA", "EMPRESA"], columns="Masa", values="CONTEO").fillna(0).reset_index()
    cilindros_lb.columns.name = None
    columnas_fijas = ["ID_EMPRESA", "EMPRESA"]
    columnas_masa = sorted([c for c in cilindros_lb.columns if c not in columnas_fijas], key=lambda x: float(x))
    cilindros_lb = cilindros_lb[columnas_fijas + columnas_masa]
    cilindros_lb["Total_Cilindros"] = cilindros_lb[columnas_masa].sum(axis=1)
    cilindros_lb["Capacidad_LB"] = round(sum(cilindros_lb[col] * float(col) for col in columnas_masa), 0)
    cilindros_lb["Cap1 i,t (kg)"] = round(cilindros_lb["Capacidad_LB"] * 0.454 * 6, 0)
    cilindros_lb.rename(columns={"ID_EMPRESA": "Código SUI", "EMPRESA": "Empresa"}, inplace=True)

    # --- 4. Cilindros en KG ---
    Cilindros_KG = ( Cilindros[Cilindros['Medida'] == 'KG'].groupby(['ID_EMPRESA', 'EMPRESA', 'Masa'], as_index=False)['CONTEO'].sum() )
    Cilindros_KG = Cilindros_KG.pivot_table(index=['ID_EMPRESA', 'EMPRESA'], columns='Masa', values='CONTEO', fill_value=0).reset_index()
    cols_num = sorted([col for col in Cilindros_KG.columns if col not in ['ID_EMPRESA','EMPRESA']], key=float)
    Cilindros_KG = Cilindros_KG[['ID_EMPRESA', 'EMPRESA'] + cols_num]
    Cilindros_KG['Total_Cilindros'] = Cilindros_KG[cols_num].sum(axis=1)
    Cilindros_KG['Capacidad_kg'] = (Cilindros_KG[cols_num] * list(map(int, cols_num))).sum(axis=1)
    Cilindros_KG['Cap2 i,t (kg)'] = (Cilindros_KG['Capacidad_kg'] * 6).round(0)
    Cilindros_KG.rename(columns={'ID_EMPRESA':'Código SUI', 'EMPRESA':'Empresa'}, inplace=True)

    # --- 5. Consolidar LB y KG ---
    data_Cilindros_CapCil = pd.merge( cilindros_lb[['Código SUI', 'Cap1 i,t (kg)']],  Cilindros_KG[['Código SUI', 'Cap2 i,t (kg)']], on='Código SUI', how='outer' )
    data_Cilindros_CapCil['Cap_cil_kg'] = data_Cilindros_CapCil[['Cap1 i,t (kg)', 'Cap2 i,t (kg)']].sum(axis=1, skipna=True)
    data_Cilindros_CapCil = data_Cilindros_CapCil.sort_values('Código SUI').reset_index(drop=True)

    if debug:
        # Devolver un diccionario con todos los resultados parciales y finales
        return {
            'Cilindros' : Cilindros,
            'cilindros_lb': cilindros_lb,
            'Cilindros_KG': Cilindros_KG,
            'data_Cilindros_CapCil': data_Cilindros_CapCil
        }
       
    return data_Cilindros_CapCil


############################################funcion para procesar capacidad RED mercado iniciales - Cap.MerIni¶ #########################


def procesar_activos_completo(activos: pd.DataFrame, proyecc_demand_usuar: pd.DataFrame, debug: bool = False) -> pd.DataFrame:
    # --- TANQUES ---
    tanques = activos.loc[activos["tipo_activo"] == "Tanques",  ["NUMERO_SOLICITUD", "RAZON_SOCIAL", "MUNICIPIO", "NOMBRE_UC", "CANTIDAD_ANO1", "RESOLUCION", "MERCADO", "ID_MERCADO", "tipo_activo", "ID_SUI"]].copy()
    
    # Extraer galones de NOMBRE_UC y calcular total
    tanques["Galones"] = pd.to_numeric(tanques["NOMBRE_UC"].str.extract(r'(\d+)')[0], errors='coerce').fillna(0).astype(int)
    tanques["cantidad_1_ajust"] = tanques["CANTIDAD_ANO1"].astype(int)
    tanques["Galones total"] = (tanques["Galones"] * tanques["cantidad_1_ajust"]).round(2)
    
    resumen_tanques = tanques.groupby( ["NUMERO_SOLICITUD", "RAZON_SOCIAL", "ID_SUI", "RESOLUCION", "MERCADO", "ID_MERCADO"] )["Galones total"].sum().reset_index()

    # --- REDES ---
    redes = activos.loc[activos["tipo_activo"] == "Redes",  ["NUMERO_SOLICITUD", "RAZON_SOCIAL", "MUNICIPIO", "NOMBRE_UC",  "CANTIDAD_ANO1", "RESOLUCION", "MERCADO", "ID_MERCADO", "tipo_activo", "ID_SUI"]].copy()
    
    redes["diametro"] = redes["NOMBRE_UC"].str.extract(r'(\d+(?:/\d+)?)')[0].apply(lambda x: float(Fraction(x)) if pd.notnull(x) else np.nan)
    redes["diametro_m"] = redes["diametro"] * 0.0254
    redes["longitud_m"] = redes["CANTIDAD_ANO1"] * 1000
    redes["volumen_m3"] = (np.pi * (redes["diametro_m"] / 2) ** 2 * redes["longitud_m"]).round(2)
    
    resumen_redes = redes.groupby( ["NUMERO_SOLICITUD", "RAZON_SOCIAL", "ID_SUI", "RESOLUCION", "MERCADO", "ID_MERCADO"] )["volumen_m3"].sum().reset_index()

    # --- PROYECCIÓN DE DEMANDA ---
    proyeccion = proyecc_demand_usuar[["NUMERO_SOLICITUD","RAZON_SOCIAL","ID_SUI","DEMANDA_1",  "RESOLUCION","MERCADO","ID_MERCADO"]].copy()
    resumen_proyeccion = proyeccion.groupby( ["NUMERO_SOLICITUD", "RAZON_SOCIAL", "ID_SUI", "RESOLUCION", "MERCADO", "ID_MERCADO"] )["DEMANDA_1"].sum().reset_index()

    # --- MERGE FINAL ---
    resumen_final = (resumen_tanques
                     .merge(resumen_redes, on=["NUMERO_SOLICITUD", "RAZON_SOCIAL", "ID_SUI",  "RESOLUCION", "MERCADO", "ID_MERCADO"], how="outer")
                     .merge(resumen_proyeccion, on=["NUMERO_SOLICITUD", "RAZON_SOCIAL", "ID_SUI", "RESOLUCION", "MERCADO", "ID_MERCADO"], how="outer"))

    # --- CONSOLIDAR ACTIVOS POR MERCADO ---
    resumen_tanques_ajust = resumen_final[resumen_final['ID_SUI'] != 0].copy()
    Llen_Ini_Est = resumen_tanques_ajust.groupby( ["RAZON_SOCIAL", "ID_SUI"] )["Galones total"].sum().reset_index()

    Llen_ini_Lin = resumen_final.groupby("ID_SUI")["volumen_m3"].sum().reset_index()
    Dem_Proy = resumen_final.groupby("ID_SUI")["DEMANDA_1"].sum().reset_index()

    base_ini = (Llen_Ini_Est.merge(Llen_ini_Lin, on='ID_SUI', how='outer')
                            .merge(Dem_Proy, on='ID_SUI', how='outer'))

    base_ini["Cap_Mer_Ini"] = ( (0.85 * base_ini["Galones total"]) * 2.01 + (base_ini["volumen_m3"] + 0.5 * base_ini["DEMANDA_1"]) * 2.11 ).round(2)
    
    #ajustar los nombre de las columnas
    base_ini = base_ini.rename(columns={
    'Galones total': 'LlenIniEst',
    'volumen_m3': 'LlenIniLin',
    'DEMANDA_1': 'DemProy'
      })
  
    if debug:
        return {
            "tanques": resumen_tanques,
            "redes": resumen_redes,
            "proyeccion": resumen_proyeccion,
            "resumen_final": resumen_final,
            "consolidado": base_ini
        }

    return base_ini


##########################################    funcion para procesar capacidad RED mercado operacion - Cap.MerOpe¶ ###############################


### para calcular capacidad de GLP por redes para mercados en operacion

def calcular_indice_mensual_continuo(usuarios_redes: pd.DataFrame, inicio_periodo: int) -> pd.DataFrame:
    # =========================
    # Índice mensual continuo
    # =========================
    mes_total = usuarios_redes['ANIO'] * 12 + usuarios_redes['PERIODO']
    mes_total_ajustado = mes_total - inicio_periodo
    
    bloque_6m = mes_total_ajustado // 6
    max_bloque = bloque_6m.max()
    
    semestre_num = (max_bloque - bloque_6m) + 1
    usuarios_redes['SEMESTRE'] = 'P' + semestre_num.astype(int).astype(str)
    
    usuarios_redes['PERIODO_MES'] = (mes_total_ajustado % 6) + 1
    
    # Ordenar y crear SEMESTRE_NUM
    df = usuarios_redes.sort_values(['ID_EMPRESA', 'Prestador',  'ID_MERCADO', 'ANIO', 'PERIODO']).copy()
    df['SEMESTRE_NUM'] = df['SEMESTRE'].str[1:].astype(int) - 1
    
    # =========================
    # Validación de semestres
    # =========================
    def validar_grupo(g):
        t = g['SEMESTRE_NUM'].max()
        ventana_A = g[((g['SEMESTRE_NUM'] == t-1) & (g['PERIODO_MES'] >= 4)) | ((g['SEMESTRE_NUM'] == t-2) & (g['PERIODO_MES'] <= 3))]
        ventana_B = g[((g['SEMESTRE_NUM'] == t-2) & (g['PERIODO_MES'] >= 4)) | ((g['SEMESTRE_NUM'] == t-1) & (g['PERIODO_MES'] <= 3))]
        return pd.Series({
            'valid_A': (ventana_A['SUSCRIPTORES'] > 0).all(),
            'valid_B': (ventana_B['SUSCRIPTORES'] > 0).all()
        })
    
    cols = ['ID_EMPRESA', 'ID_MERCADO', 'Prestador']
    df2 = df[cols + ['SEMESTRE_NUM', 'PERIODO_MES', 'SUSCRIPTORES']].copy()
    
    resultado = df2.groupby(cols).apply(validar_grupo).reset_index()
    resultado['valida_total'] = resultado['valid_A'] & resultado['valid_B']
    
    # =========================
    # Promedios de consumo y usuarios
    # =========================
    t = df['SEMESTRE_NUM'].min()
    t_1, t_2 = t + 1, t + 2
    
    # Función auxiliar para calcular promedio de una columna
    def promedio_por_semestre(df, semestre, meses, columna, nombre_col):
        subset = df[(df['SEMESTRE_NUM'] == semestre) & (df['PERIODO_MES'].isin(meses))]
        promedio = subset.groupby(['ID_EMPRESA', 'ID_MERCADO', 'Prestador'])[columna].mean().round(0).reset_index()
        promedio.rename(columns={columna: nombre_col}, inplace=True)
        return promedio
    
    promedio_consumo_t1 = promedio_por_semestre(df, t_1, [1,2,3], 'CONSUMO', 'PROM_CONSUMO_T1_M_3')
    promedio_usuarios_t1 = promedio_por_semestre(df, t_1, [1,2,3], 'SUSCRIPTORES', 'PROM_USUARIOS_T1')
    promedio_usuarios_t2 = promedio_por_semestre(df, t_2, [1,2,3], 'SUSCRIPTORES', 'PROM_USUARIOS_T2')
    
    # =========================
    # Unión de todas las tablas
    # =========================
    base_final = (promedio_usuarios_t1
                  .merge(promedio_usuarios_t2, on=['ID_EMPRESA', 'ID_MERCADO','Prestador'], how='outer')
                  .merge(promedio_consumo_t1, on=['ID_EMPRESA', 'ID_MERCADO','Prestador'], how='outer')
                  .merge(resultado, on=['ID_EMPRESA', 'ID_MERCADO','Prestador'], how='outer'))
    
    
    # =========================
    # Calcular Cap_Mer_Opera
    # =========================
    base_final["Cap_Mer_Opera"] = ( base_final["PROM_CONSUMO_T1_M_3"] * 2.11 * (base_final["PROM_USUARIOS_T1"] / base_final["PROM_USUARIOS_T2"]) * 6 ).round(2)
    
     #ajustar los nombre de las columnas
    base_final = base_final.rename(columns={
    'PROM_CONSUMO_T1_M_3': 'Vt_1',
    'PROM_USUARIOS_T1': 'NSt_1',
    'PROM_USUARIOS_T2': 'NSt_2'
      })
            
    return base_final


####################################################################  funcion para procesar capacidad RED - Cap.red ###################################################
### para consolidar la capacidad de GLP por redes para mercados en operacion
def consolidar_activos(resumen_activos: pd.DataFrame, base_final: pd.DataFrame) -> pd.DataFrame:
    # --- Tanques ---
    resumen_tanques_ajust = resumen_activos[resumen_activos['ID_SUI'] != 0].copy()
    Llen_Ini_Est = resumen_tanques_ajust.groupby( ["RAZON_SOCIAL", "ID_SUI"] )["LlenIniEst"].sum().reset_index()

    # --- Redes ---
    Llen_ini_Lin = resumen_activos.groupby("ID_SUI")["LlenIniLin"].sum().reset_index()

    # --- Proyección de demanda ---
    Dem_Proy = resumen_activos.groupby("ID_SUI")["DemProy"].sum().reset_index()

    # --- Merge inicial ---
    base_ini = (Llen_Ini_Est.merge(Llen_ini_Lin, on='ID_SUI', how='outer')
                            .merge(Dem_Proy, on='ID_SUI', how='outer'))

    # --- Calcular Cap_Mer_Ini ---
    base_ini["Cap_Mer_Ini"] = ( (0.85 * base_ini["LlenIniEst"]) * 2.01 + (base_ini["LlenIniLin"] + 0.5 * base_ini["DemProy"]) * 2.11 ).round(2)

    # --- Calcular Cap_Mer_Opera por empresa y prestador ---
    Cap_Mer_Opera = base_final.groupby(["ID_EMPRESA", "Prestador"])["Cap_Mer_Opera"].sum().reset_index()

    # --- Merge con base_ini para obtener Cap_red ---
    cap_red = pd.merge( base_ini,  Cap_Mer_Opera,  left_on='ID_SUI',  right_on='ID_EMPRESA', how='inner' )

    cap_red["Cap_red"] = (cap_red["Cap_Mer_Ini"] + cap_red["Cap_Mer_Opera"]).round(2)
    
    cap_red = cap_red.drop(columns=['ID_EMPRESA', 'Prestador'])

    return cap_red


##########################################    funcion para consolidar todas las capacidades cilindros, tanques y redes#############################


def consolidar_capacidades(data_Tanques_CapTE_final: pd.DataFrame,
                            data_Cilindros_CapCil: pd.DataFrame,
                            resumen_activos: pd.DataFrame,
                            base_final: pd.DataFrame,
                            debug: bool = False) -> pd.DataFrame:
    
    # --- 1. Merge de tanques y cilindros ---
    merged_data = pd.merge( data_Tanques_CapTE_final,  data_Cilindros_CapCil,  on='Código SUI',  how='outer' )

    # --- 2. Reemplazar NA por 0 y calcular CCit ---
    merged_data['CapTEi_t_kg'] = merged_data['CapTEi_t_kg'].fillna(0)
    merged_data['Cap_cil_kg'] = merged_data['Cap_cil_kg'].fillna(0)

    # --- 3. Consolidar capacidades de redes usando resumen_completo ---
    # Suponiendo que consolidar_activos() ya devuelve Cap_Mer_Ini por ID_SUI
    Capacidad_Redes = consolidar_activos(resumen_activos,base_final)

    # Hacer merge con merged_data para incorporar redes)
    merged_data = pd.merge(  merged_data,  Capacidad_Redes[['ID_SUI', 'Cap_red']],  left_on='Código SUI',  right_on='ID_SUI',  how='left' )

    merged_data['Cap_red'] = merged_data['Cap_red'].fillna(0)
    
    merged_data['Tiene_Cilindros'] = (merged_data['Cap_cil_kg'] > 0).astype(int)
    

    # --- 4. Calcular CCit final considerando tanques, cilindros y redes ---
    # Fórmula: CCit_kg = (0.85 * CapTEi_t_kg + Cap_cil_kg + Cap_Mer_Ini) * 0.345
    
    merged_data['CCit_kg'] = ((0.85 * merged_data['CapTEi_t_kg'] +  merged_data['Cap_cil_kg'] ) * 0.345 *  merged_data['Tiene_Cilindros'] + merged_data['Cap_red'] ).round(0)

    # --- 5. Seleccionar columnas finales y ordenar ---
    merged_data_final = merged_data[['Código SUI', 'CapTEi_t_kg', 'Cap_cil_kg', 'Cap_red', 'CCit_kg']]
    merged_data_final = merged_data_final.sort_values('Código SUI').reset_index(drop=True)

    if debug:
        return {
            "merged_data": merged_data,
            "merged_data_final": merged_data_final,
            "Capacidad_Redes": Capacidad_Redes
        }

    return merged_data_final


###############################################   Funcion para limpieza de nombre #######################################################



## Funciones de limpieza de los nombre  y codigo SUI
def limpiar_id(x):
    return str(x).strip().lstrip('0')

def limpiar_empresa(x):
    x_str = str(x).strip()
    x_str = re.sub(r'\s*\d{3,}$', '', x_str)
    return x_str

# Normalización de nombres
def normalizar_nombre_empresa(nombre):
    nombre = str(nombre)
    # Reemplazar "EMPRESA DE SERVICIOS PÚBLICOS" o "DOMICILIARIOS" → "E.S.P."
    nombre = re.sub(
        r"EMPRESA DE SERVICIOS P[ÚU]BLICOS( DOMICILIARIOS)?",
        "E.S.P.",
        nombre,
        flags=re.IGNORECASE
    )
    # Reemplazar "SOCIEDAD POR ACCIONES SIMPLIFICADA" → "S.A.S."
    nombre = re.sub(  r"SOCIEDAD POR ACCIONES SIMPLIFICADA", "S.A.S.",   nombre,  flags=re.IGNORECASE )
    # Eliminar espacios dobles y sobrantes
    nombre = re.sub(r"\s+", " ", nombre).strip()

    return nombre


def consolidar_agentes(base_empresas_df: pd.DataFrame, #Capacidad_GLP["merged_data_final"]
                       Cilindros: pd.DataFrame,
                       resultados_tan: dict,
                       Cap_Mer_Ini: pd.DataFrame,
                       Cap_Mer_Ope: pd.DataFrame,
                       limpiar_id,
                       limpiar_empresa,
                       debug: bool = False) -> pd.DataFrame:
               
    # --- Base de empresas ---
    base_empresas = base_empresas_df[['Código SUI']].drop_duplicates().copy()
    base_empresas['Código SUI'] = base_empresas['Código SUI'].apply(limpiar_id)
    
    # --- Cilindros por empresa ---
    Cilindros_por_empresa = (Cilindros.groupby(['ID_EMPRESA', 'EMPRESA', 'Medida'], as_index=False).agg(Total_Cilindros=('CONTEO', 'sum')))
    
    cil_unico = (Cilindros_por_empresa[['ID_EMPRESA', 'EMPRESA']].drop_duplicates(subset='ID_EMPRESA').copy())
    
    cil_unico['ID_EMPRESA'] = cil_unico['ID_EMPRESA'].apply(limpiar_id)
    cil_unico['EMPRESA'] = cil_unico['EMPRESA'].apply(limpiar_empresa)
    
    # --- Tanques T2 ---
    tanq_T2_unico = (resultados_tan['T2_clean'][['ID_EMPRESA', 'EMPRESA']].drop_duplicates(subset='ID_EMPRESA').copy())
    tanq_T2_unico['ID_EMPRESA'] = tanq_T2_unico['ID_EMPRESA'].apply(limpiar_id)
    tanq_T2_unico['EMPRESA'] = tanq_T2_unico['EMPRESA'].apply(limpiar_empresa)
    
    # --- Tanques T1 ---
    tanq_T1_unico = (resultados_tan['T1_clean'][['ID_EMPRESA', 'EMPRESA']].drop_duplicates(subset='ID_EMPRESA').copy())
    tanq_T1_unico['ID_EMPRESA'] = tanq_T1_unico['ID_EMPRESA'].apply(limpiar_id)
    tanq_T1_unico['EMPRESA'] = tanq_T1_unico['EMPRESA'].apply(limpiar_empresa)
    
    # --- Red inicial ---
    glp_Red_inic_unic = (Cap_Mer_Ini[['ID_SUI', 'RAZON_SOCIAL']].drop_duplicates(subset='ID_SUI').copy())
    glp_Red_inic_unic['ID_SUI'] = glp_Red_inic_unic['ID_SUI'].apply(limpiar_id)
    glp_Red_inic_unic['RAZON_SOCIAL'] = glp_Red_inic_unic['RAZON_SOCIAL'].apply(limpiar_empresa)
    
    # --- Red operativa ---
    glp_Red_opera_unico = (Cap_Mer_Ope[['ID_EMPRESA', 'Prestador']].drop_duplicates(subset='ID_EMPRESA').copy())
    glp_Red_opera_unico['ID_EMPRESA'] = glp_Red_opera_unico['ID_EMPRESA'].apply(limpiar_id)
    glp_Red_opera_unico['Prestador'] = glp_Red_opera_unico['Prestador'].apply(limpiar_empresa)
    
    # --- Merge Cilindros ---
    agentes_data = base_empresas.merge(cil_unico, left_on='Código SUI', right_on='ID_EMPRESA', how='left' )
    
    # --- Merge Tanques T2 ---
    agentes_data = agentes_data.merge( tanq_T2_unico, left_on='Código SUI', right_on='ID_EMPRESA', how='left', suffixes=('', '_T2'))
    
    # --- Merge Tanques T1 ---
    agentes_data = agentes_data.merge(tanq_T1_unico, left_on='Código SUI', right_on='ID_EMPRESA', how='left', suffixes=('', '_T1'))
    
   # --- Merge Red inici ---
    agentes_data= agentes_data.merge( glp_Red_inic_unic, left_on='Código SUI', right_on='ID_SUI', how='left', suffixes=('', '_Red_ini'))
    
    # --- Merge Red opera ---
    agentes_data = agentes_data.merge( glp_Red_opera_unico, left_on='Código SUI', right_on='ID_EMPRESA', how='left', suffixes=('', '_Red_Ope'))
       
        
    # --- Coalesce de nombres de empresa ---
    agentes_data['EMPRESA'] = agentes_data['EMPRESA'].combine_first(agentes_data['EMPRESA_T2']
                             ).combine_first( agentes_data['EMPRESA_T1']
                             ).combine_first( agentes_data.get('RAZON_SOCIAL')  # Red inicial
                             ).combine_first( agentes_data.get('Prestador')      # Red operador
                             )
    
    # --- Eliminar columnas auxiliares ---
    drop_cols = ['EMPRESA_T2', 'EMPRESA_T1', 'ID_EMPRESA', 'ID_EMPRESA_T1', 'ID_EMPRESA_T2', 'ID_SUI', 'RAZON_SOCIAL', 'Prestador', 'ID_EMPRESA_Red_Ope']
    agentes_data = agentes_data.drop(columns=drop_cols, errors='ignore')
 
    
    if debug:
        return {
            "base_empresas": base_empresas,
            "cil_unico": cil_unico,
            "tanq_T2_unico": tanq_T2_unico,
            "tanq_T1_unico": tanq_T1_unico,
            "glp_Red_inic_unic": glp_Red_inic_unic,
            "glp_Red_opera_unico": glp_Red_opera_unico,
            "agentes_data": agentes_data
        }
    
    return agentes_data



def validar_agentes(agentes_df: pd.DataFrame) -> None:
   
    # --- 1. Verificar duplicados ---
    duplicados = agentes_df['Código SUI'].duplicated().any()
    if duplicados:
        duplicados_ids = agentes_df[agentes_df['Código SUI'].duplicated(keep=False)]['Código SUI'].unique()
        raise ValueError(f"ERROR: existen IDs duplicados: {', '.join(map(str, duplicados_ids))}")

    # --- 2. Verificar nombres de empresa ---
    ids_sin_nombre = agentes_df[ agentes_df['EMPRESA'].isna() | (agentes_df['EMPRESA'].str.strip() == '') ]['Código SUI'].tolist()

    if len(ids_sin_nombre) > 0:
        raise ValueError(
            f"ERROR: existen {len(ids_sin_nombre)} empresas sin nombre. IDs afectados: {', '.join(map(str, ids_sin_nombre))}"
        )

    print("Validaciones completadas: no hay IDs duplicados y todas las empresas tienen nombre.")



#####################################################  para visualización¶ ##########################################################



def crear_indice_empresas(agentes_df: pd.DataFrame,
                           merged_data: pd.DataFrame,
                           normalizar_nombre_empresa) -> pd.DataFrame:
      
    df = agentes_df.copy()
    
    # Asegurar que 'Código SUI' sea string
    df['Código SUI'] = df['Código SUI'].astype(str)
    
    # Ordenar por Código SUI y asignar índice
    df = df.sort_values('Código SUI').reset_index(drop=True)
    df['indice_empresa'] = range(1, len(df) + 1)
    
    # Seleccionar columnas clave
    df_final = df[['indice_empresa', 'Código SUI', 'EMPRESA']].copy()
    
    # Normalizar nombres de empresa
    df_final['EMPRESA'] = df_final['EMPRESA'].apply(normalizar_nombre_empresa)
    
    # Asegurar que Código SUI sea numérico para merge
    df['Código_SUI_num'] = pd.to_numeric(df['Código SUI'], errors='coerce')
    merged_data['Código_SUI_num'] = pd.to_numeric(merged_data['Código SUI'], errors='coerce')
    
    
    # Merge con capacidades CCit_kg
    tabla_final = df.merge( merged_data[['Código_SUI_num', 'CCit_kg']], on='Código_SUI_num',  how='left' )
    
      # Seleccionar columnas finales
    tabla_final = tabla_final[['indice_empresa', 'Código SUI', 'EMPRESA', 'CCit_kg']]
    
    
    # ============================================================
    # Calcular participación de cada empresa respecto al total
    # ============================================================
    
    total_CCit_kg = tabla_final['CCit_kg'].sum(skipna=True)  # suma total
    tabla_final['participacion'] = tabla_final['CCit_kg'] / total_CCit_kg
    
    return tabla_final



def visualizar_y_tabla_final(agentes_df: pd.DataFrame, merged_data: pd.DataFrame, merged_data_Red: pd.DataFrame ) -> tuple[pd.DataFrame, float]:
      
    df_agentes = agentes_df.copy()
    df_merged = merged_data.copy()
    
    # --- 7.1 Estadísticas ---
    CC_total = df_merged['CCit_kg'].sum(skipna=True)
    
    cil = df_merged['Cap_cil_kg'].sum(skipna=True)
    tan = df_merged['CapTEi_t_kg'].sum(skipna=True)
    
    red_gener = df_merged['Cap_red'].sum(skipna=True)
    
    red = merged_data_Red['Cap_red'].sum(skipna=True)
    red_ini = merged_data_Red['Cap_Mer_Ini'].sum(skipna=True)
    red_ope = merged_data_Red['Cap_Mer_Opera'].sum(skipna=True)
       
    tan_cap = round((0.85 * tan + 0) * 0.345, 0)

    return {
        "Numero_empresas" : df_merged['Código SUI'].nunique(),
        "CC_total" :  CC_total,
        "prop_cilindros_CC_total":  (cil / CC_total if CC_total != 0 else 0),
        "prop_tanques_CC_total": (tan_cap / CC_total if CC_total != 0 else 0),
        "prop_Red_CC_total": (red_gener / CC_total if CC_total != 0 else 0),
        "prop_Red ini_Red":  (red_ini / red if CC_total != 0 else 0),
        "prop_Red_opera_Red": (red_ope / red if CC_total != 0 else 0)
        } 
    return CC_total


def resumen_tanques_cilindros(data_tanques_T1_clean, data_tanques_T2_clean, Cilindros):
    # ============================================================
    # 1. INFORMACIÓN DE TANQUES ESTACIONARIOS
    # ============================================================
    
    # Trim 3 (T1)
    empresas_tanques_T3 = data_tanques_T1_clean[data_tanques_T1_clean['CapTE_T1'] > 0]
    n_empresas_tanques_T3 = empresas_tanques_T3['ID_EMPRESA'].nunique()
    n_tanques_T3 = empresas_tanques_T3['n_T1'].sum(skipna=True)
    
    # Trim 4 (T2)
    empresas_tanques_T4 = data_tanques_T2_clean[data_tanques_T2_clean['CapTE_T2'] > 0]
    n_empresas_tanques_T4 = empresas_tanques_T4['ID_EMPRESA'].nunique()
    n_tanques_T4 = empresas_tanques_T4['n_T2'].sum(skipna=True)
    
    # ============================================================
    # 2. INFORMACIÓN DE CILINDROS
    # ============================================================
    
    empresas_cilindros = Cilindros[Cilindros['CONTEO'] > 0]
    n_empresas_cilindros = empresas_cilindros['ID_EMPRESA'].nunique()
    n_marcas_cilindros = empresas_cilindros['MARCA'].nunique()
    n_cilindros_total = Cilindros['CONTEO'].sum(skipna=True)
    
    # ============================================================
    # Resultado en diccionario
    # ============================================================
    resumen = {
        "empresas_tanques_T3": n_empresas_tanques_T3,
        "tanques_T3": n_tanques_T3,
        "empresas_tanques_T4": n_empresas_tanques_T4,
        "tanques_T4": n_tanques_T4,
        "empresas_cilindros": n_empresas_cilindros,
        "marcas_cilindros": n_marcas_cilindros,
        "cilindros_total": n_cilindros_total
    }
    
    return resumen


def crear_resumen_presentacion(resumen_dict):
    indicadores = [
        "Empresas operativas con tanques Trim 3",
        "Cantidad de tanques Trim 3",
        "Empresas operativas con tanques Trim 4",
        "Cantidad de tanques Trim 4",
        "Empresas operativas con cilindros",
        "Cantidad de marcas con registros",
        "Cantidad total de cilindros"
    ]
    
    valores = [
        resumen_dict["empresas_tanques_T3"],
        resumen_dict["tanques_T3"],
        resumen_dict["empresas_tanques_T4"],
        resumen_dict["tanques_T4"],
        resumen_dict["empresas_cilindros"],
        resumen_dict["marcas_cilindros"],
        resumen_dict["cilindros_total"]
    ]
    
    resumen_estadistico = pd.DataFrame({
        "Indicador": indicadores,
        "Valor": valores
    })
    
    return resumen_estadistico


######################################### funcion para pasar a word ######################################
def generar_word(df, titulo="Reporte"):

    doc = Document()

    doc.add_heading(titulo, level=1)

    tabla = doc.add_table(rows=1, cols=len(df.columns))
    tabla.style = 'Table Grid'

    # encabezados
    hdr = tabla.rows[0].cells

    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)

    # datos
    for _, row in df.iterrows():

        cells = tabla.add_row().cells

        for i, val in enumerate(row):

            if isinstance(val, (int, float)):
                cells[i].text = f"{val:,.0f}"
            else:
                cells[i].text = str(val)

    buffer = BytesIO()

    doc.save(buffer)

    buffer.seek(0)

    return buffer


def agregar_tabla(doc, df, titulo):

    # =========================
    # TITULO
    # =========================
    titulo_doc = doc.add_heading(titulo, level=2)
    titulo_doc.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # =========================
    # TAMAÑO REAL DE PÁGINA
    # =========================
    section = doc.sections[0]

    page_width = section.page_width
    left_margin = section.left_margin
    right_margin = section.right_margin

    usable_width = page_width - left_margin - right_margin

    # =========================
    # TABLA
    # =========================
    tabla = doc.add_table(rows=1, cols=len(df.columns))
    tabla.style = 'Table Grid'
    tabla.alignment = WD_TABLE_ALIGNMENT.LEFT
    tabla.autofit = False

    # ancho por columna (REAL dentro de la hoja)
    col_width = usable_width / len(df.columns)

    # =========================
    # ENCABEZADOS
    # =========================
    hdr = tabla.rows[0].cells

    for i, col in enumerate(df.columns):

        hdr[i].width = col_width

        p = hdr[i].paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        sombrear_celda(hdr[i], color="BFBFBF")  # gris más oscuro

        run = p.add_run(str(col))
        run.bold = True
        run.font.size = Pt(7)

    # =========================
    # FILAS
    # =========================
    for _, row in df.iterrows():

        cells = tabla.add_row().cells

        for i, val in enumerate(row):

            texto = f"{val:,.0f}" if isinstance(val, (int, float)) else str(val)

            cells[i].width = col_width

            p = cells[i].paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1

            run = p.add_run(texto)
            run.font.size = Pt(7)

            if isinstance(val, (int, float)):
                p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            else:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph("")

def generar_word_completo(tablas):

    doc = Document()

    # 🔥 AQUÍ va la configuración global
    section = doc.sections[0]

    section.orientation = WD_ORIENT.LANDSCAPE

    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height

    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    for titulo, df in tablas:
        agregar_tabla(doc, df, titulo)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer

def calcular_ancho_columnas_proporcional(df, max_width_inches=10):

    cols = len(df.columns)

    # pesos base: encabezado + contenido
    pesos = []

    for col in df.columns:

        max_len = len(str(col))

        for val in df[col]:
            max_len = max(max_len, len(str(val)))

        pesos.append(max_len)

    total = sum(pesos)

    # evitar división por cero
    if total == 0:
        return [Inches(max_width_inches / cols)] * cols

    anchos = []

    for p in pesos:

        ancho = (p / total) * max_width_inches

        # límites para que no se rompa el diseño
        ancho = max(0.8, min(ancho, 3))

        anchos.append(Inches(ancho))

    return anchos

def sombrear_celda(cell, color="D9D9D9"):

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)

    tcPr.append(shd)


 ######################################### funcion para pasar comparar con el periodo anterior ######################################
    
def comparar_workbooks(actual, anterior, key="id_empresa", columnas_ignorar=None):

    if columnas_ignorar is None:
        columnas_ignorar = []

    resultados = {}

    hojas = actual.keys() & anterior.keys()

    for hoja in hojas:

        df_act = actual[hoja].copy()
        df_ant = anterior[hoja].copy()

        df_act.columns = [str(c).strip() for c in df_act.columns]
        df_ant.columns = [str(c).strip() for c in df_ant.columns ]

        # -------------------------
        # llave segura
        # -------------------------
        df_act[key] = df_act[key].astype(str).str.strip()
        df_ant[key] = df_ant[key].astype(str).str.strip()

        # -------------------------
        # merge
        # -------------------------
        df = df_act.merge(
            df_ant,
            on=key,
            suffixes=("_act", "_ant"),
            how="inner"
        )

        resultado = pd.DataFrame()
        resultado[key] = df[key]

        columnas_base = sorted(set(
            c.replace("_act", "")
            for c in df.columns
            if c.endswith("_act")
        ))

        for col in columnas_base:


            if col in columnas_ignorar:
                continue

            col_act = f"{col}_act"
            col_ant = f"{col}_ant"

            if col_act not in df.columns or col_ant not in df.columns:
                continue

            act_raw = df[col_act]
            ant_raw = df[col_ant]

            # =====================================================
            # 🔥 CONVERSIÓN SEGURA REAL (CLAVE DEL FIX)
            # =====================================================

            act_num = pd.to_numeric(act_raw, errors="coerce")
            ant_num = pd.to_numeric(ant_raw, errors="coerce")

            # condición REAL: si al menos uno tiene números válidos
            es_numerico = (
                act_num.notna().any() or
                ant_num.notna().any()
            )

           # evitar columnas booleanas
            es_bool = (
                pd.api.types.is_bool_dtype(act_raw) or
                pd.api.types.is_bool_dtype(ant_raw)
            )

            if es_numerico and not es_bool:
                  
                try:
                    diff = act_num - ant_num

                    resultado[f"{col}_actual"] = act_num
                    resultado[f"{col}_anterior"] = ant_num

                    resultado[f"diff_{col}"] = diff
                    resultado[f"pct_{col}"] = diff / ant_num.replace(0, pd.NA)
                except Exception:
                    # si falla la resta, tratar como texto
                    resultado[f"{col}_actual"] = act_raw
                    resultado[f"{col}_anterior"] = ant_raw
                    resultado[f"diff_{col}"] = None
                    resultado[f"pct_{col}"] = None  

            else:

                # 👇 aquí nunca se hace resta
                resultado[f"{col}_actual"] = act_raw
                resultado[f"{col}_anterior"] = ant_raw
                resultado[f"diff_{col}"] = None
                resultado[f"pct_{col}"] = None

        resultados[hoja] = resultado

    return resultados


def estandarizar_llave(df):

    df = df.copy()

    # intenta encontrar cualquier columna posible
    posibles = ["Código SUI", "ID_SUI", "ID_EMPRESA", "Código SUI / Capacidad"]

    for col in posibles:
        if col in df.columns:
            df["id_empresa"] = df[col].astype(str)
            break

    return df


####################################################### funcion consolidando dos funciones generales ###########################


## Para cilindros y tanques ###########
def ejecutar_cilindros_tanques(
    df_cilindros,
    df_tanques_TM,
    df_tanques_TM_1
):

    # 1. Cálculo general
    General = calcular_CCit(
        df_cilindros,
        df_tanques_TM,
        df_tanques_TM_1
    )

    # 2. Cilindros
    resultados_cil = procesar_cilindros(
        df_cilindros,
        True
    )

    # 3. Tanques
    resultados_tan = procesar_tanques(
        df_tanques_TM,
        df_tanques_TM_1,
        True
    )

    # -----------------------------
    # Tablas resolución
    # -----------------------------

    Cilindros_LB = (
        resultados_cil["cilindros_lb"]
        .rename(columns={
            "Código SUI": "Código SUI / Capacidad",
            "Total_Cilindros": "Total Cilindros",
            "Capacidad_LB": "Capacidad (Libras)"
        })
    )

    Cilindros_KG = (
        resultados_cil["Cilindros_KG"]
        .rename(columns={
            "Código SUI": "Código SUI / Capacidad",
            "Total_Cilindros": "Total general",
            "Capacidad_kg": "Capacidad (kg)"
        })
    )

    Cilindros_consol = (
        resultados_cil["data_Cilindros_CapCil"]
        .rename(columns={
            "Cap_cil_kg": "Cap. cil (kg)"
        })
    )

    resultados_tanques = (
        resultados_tan["Final"]
        .rename(columns={
            "Capacidad_gal": "Capacidad (gal)",
            "Número_de_Tanques": "Número de Tanques",
            "CapTEi_t_kg": "CapTEi,t (kg)"
        })
    )

    resultado_CC_tanques = (
        General["resultado"]
        .rename(columns={
            "CapTE": "CapTEi,t (kg)",
            "Cap_cil": "Cap.cil (kg)",
            "CCit": "(CapTE * 0.85 + Cap_cil) * 0.345"
        })
    )

    resumen = resumen_tanques_cilindros(
        resultados_tan["T1_clean"],
        resultados_tan["T2_clean"],
        resultados_cil["Cilindros"]
    )

    # ==========================
    # AQUÍ va num_empresas
    # ==========================

    empresas_t1 = set(
        resultados_tan["T1_clean"]
        .loc[resultados_tan["T1_clean"]["CapTE_T1"] > 0, "ID_EMPRESA"]
    )

    empresas_t2 = set(
        resultados_tan["T2_clean"]
        .loc[resultados_tan["T2_clean"]["CapTE_T2"] > 0, "ID_EMPRESA"]
    )

    empresas_cil = set(
        resultados_cil["Cilindros"]
        .loc[resultados_cil["Cilindros"]["CONTEO"] > 0, "ID_EMPRESA"]
    )

    num_empresas = len(empresas_t1 | empresas_t2 | empresas_cil)
    Capacidad_GLP_tot = resultados_tan["Final"]

    return {
        "General": General,
        "resultados_cil": resultados_cil,
        "resultados_tan": resultados_tan,
        "Cilindros_LB": Cilindros_LB,
        "Cilindros_KG": Cilindros_KG,
        "Cilindros_consol": Cilindros_consol,
        "resultados_tanques": resultados_tanques,
        "resultado_CC_tanques": resultado_CC_tanques,

        # Nuevos indicadores
        "resumen": resumen,
        "empresas_cilindros": resumen["empresas_cilindros"],
        "marcas_cilindros": resumen["marcas_cilindros"],
        "cilindros_total": resumen["cilindros_total"],
        "empresas_tanques_T3": resumen["empresas_tanques_T3"],
        "tanques_T3": resumen["tanques_T3"],
        "empresas_tanques_T4": resumen["empresas_tanques_T4"],
        "tanques_T4": resumen["tanques_T4"],
        "num_empresas": num_empresas, 
        "Capacidad_GLP_tot": Capacidad_GLP_tot,
    }

## Para cilindros, tanques y redes ###########
def ejecutar_completo(
    df_cilindros,
    df_tanques_TM,
    df_tanques_TM_1,
    df_inversio_nuev,
    df_proyec_usuar_demand,
    df_red,
    inicio_periodo
):

    resultados = ejecutar_cilindros_tanques(
        df_cilindros,
        df_tanques_TM,
        df_tanques_TM_1
    )

    # -----------------------------
    # Redes
    # -----------------------------

    Cap_Mer_Ini = procesar_activos_completo(
        df_inversio_nuev,
        df_proyec_usuar_demand,
        True
    )

    Cap_Mer_Ope = calcular_indice_mensual_continuo(
        df_red,
        inicio_periodo
    )

    Cap_Mer_Opera_cons = (
        Cap_Mer_Ope
        .groupby(
            ["ID_EMPRESA", "Prestador"],
            as_index=False
        )["Cap_Mer_Opera"]
        .sum()
    )

    Capacidad_Redes = consolidar_activos(
        Cap_Mer_Ini["consolidado"],
        Cap_Mer_Ope
    )

    Capacidad_GLP = consolidar_capacidades(
        resultados["resultados_tan"]["Final"],
        resultados["resultados_cil"]["data_Cilindros_CapCil"],
        Cap_Mer_Ini["consolidado"],
        Cap_Mer_Ope,
        True
    )

    tablas = preparar_tablas_resolucion(
        resultados["General"],
        resultados["resultados_cil"],
        resultados["resultados_tan"],
        Cap_Mer_Ini,
        Capacidad_GLP
    )

    resultados.update(tablas)

    resultados.update({
        "Cap_Mer_Ini": Cap_Mer_Ini,
        "Cap_Mer_Ope": Cap_Mer_Ope,
        "Cap_Mer_Opera_cons": Cap_Mer_Opera_cons,
        "Capacidad_Redes": Capacidad_Redes,
        "Capacidad_GLP": Capacidad_GLP,
        "Cap_Mer_Ini_consol": Cap_Mer_Ini["consolidado"],
        "Capacidad_GLP_tot": Capacidad_GLP["merged_data_final"]
    })
    return resultados

def generar_reportes(
    Capacidad_GLP,
    resultados_cil,
    resultados_tan,
    Cap_Mer_Ini_consol,
    Cap_Mer_Ope
):

    empresas = consolidar_agentes(
        Capacidad_GLP["merged_data_final"],
        resultados_cil["Cilindros"],
        resultados_tan,
        Cap_Mer_Ini_consol,
        Cap_Mer_Ope,
        limpiar_id,
        limpiar_empresa,
        True
    )

    Resum_gener = visualizar_y_tabla_final(
        empresas["agentes_data"],
        Capacidad_GLP["merged_data_final"],
        Capacidad_GLP["Capacidad_Redes"]
    )

    resumen = resumen_tanques_cilindros(
        resultados_tan["T1_clean"],
        resultados_tan["T2_clean"],
        resultados_cil["Cilindros"]
    )

    participa_capacidad = crear_indice_empresas(
        empresas["agentes_data"],
        Capacidad_GLP["merged_data_final"],
        normalizar_nombre_empresa
    )

    return {
        "empresas": empresas,
        "Resum_gener": Resum_gener,
        "resumen": resumen,
        "participa_capacidad": participa_capacidad,

        "num_empresas": Resum_gener["Numero_empresas"],

        "empresas_cilindros": resumen["empresas_cilindros"],
        "marcas_cilindros": resumen["marcas_cilindros"],
        "cilindros_total": resumen["cilindros_total"],

        "empresas_tanques_T3": resumen["empresas_tanques_T3"],
        "tanques_T3": resumen["tanques_T3"],

        "empresas_tanques_T4": resumen["empresas_tanques_T4"],
        "tanques_T4": resumen["tanques_T4"], 
        "Capacidad_GLP_tot": Capacidad_GLP["merged_data_final"]
    }

def obtener_hojas_excel(archivo):
    return pd.ExcelFile(archivo).sheet_names

def leer_hoja_excel(archivo, hoja):
    return pd.read_excel(archivo, sheet_name=hoja)

def preparar_tablas_resolucion(
    General,
    resultados_cil,
    resultados_tan,
    Cap_Mer_Ini=None,
    Capacidad_GLP=None
):
    
    # Capítulo 1
    Cilindros_LB = resultados_cil['cilindros_lb'].rename(columns={
        "Código SUI": "Código SUI / Capacidad",
        "Total_Cilindros": "Total Cilindros",
        "Capacidad_LB": "Capacidad (Libras)"
    })

    # Capítulo 2
    Cilindros_KG = resultados_cil['Cilindros_KG'].rename(columns={
        "Código SUI": "Código SUI / Capacidad",
        "Total_Cilindros": "Total general",
        "Capacidad_kg": "Capacidad (kg)"
    })

    # Capítulo 3
    Cilindros_consol = resultados_cil['data_Cilindros_CapCil'].rename(columns={
        "Cap_cil_kg": "Cap. cil (kg)"
    })

    # Capítulo 4
    resultados_tanques = resultados_tan['Final'].rename(columns={
        "Capacidad_gal": "Capacidad (gal)",
        "Número_de_Tanques": "Número de Tanques",
        "CapTEi_t_kg": "CapTEi,t (kg)"
    })

    # Capítulo 5
    resultado_CC_tanques = General['resultado'].rename(columns={
        "CapTE": "CapTEi,t (kg)",
        "Cap_cil": "Cap.cil (kg)",
        "CCit": "(CapTE * 0.85 + Cap_cil) * 0.345"
    })

    salida = {
        "Cilindros_LB": Cilindros_LB,
        "Cilindros_KG": Cilindros_KG,
        "Cilindros_consol": Cilindros_consol,
        "resultados_tanques": resultados_tanques,
        "resultado_CC_tanques": resultado_CC_tanques
    }

    if Cap_Mer_Ini is not None:
        salida["Cap_Mer_Ini_consol"] = Cap_Mer_Ini["consolidado"]

    if Capacidad_GLP is not None:
        salida["Capacidad_GLP_tot"] = Capacidad_GLP["merged_data_final"]

    return salida

def preparar_comparacion(resultados):

    out = {
        "Cilindros_LB": estandarizar_llave(resultados["Cilindros_LB"]),
        "Cilindros_KG": estandarizar_llave(resultados["Cilindros_KG"]),
        "Cilindros_consol": estandarizar_llave(resultados["Cilindros_consol"]),
        "resultados_tanques": estandarizar_llave(resultados["resultados_tanques"]),
        "resultados_cilindros_tanques": estandarizar_llave(resultados["resultado_CC_tanques"]),
          }

    # opcionales (solo si existen)
    if "Cap_Mer_Ini_consol" in resultados:
        out["Cap_Mer_Ini"] = estandarizar_llave(resultados["Cap_Mer_Ini_consol"])

    if "Cap_Mer_Opera_cons" in resultados:
        out["Cap_Mer_Ope"] = estandarizar_llave(resultados["Cap_Mer_Opera_cons"])

    if "Capacidad_Redes" in resultados:
        out["Capacidad_Redes"] = estandarizar_llave(resultados["Capacidad_Redes"])

    if "Capacidad_GLP_tot" in resultados:
        out["Capacidad_GLP_tot"] = estandarizar_llave(resultados["Capacidad_GLP_tot"])


    if "participa_capacidad" in resultados:
        out["participa_capacidad"] = estandarizar_llave(resultados["participa_capacidad"])

    return out

